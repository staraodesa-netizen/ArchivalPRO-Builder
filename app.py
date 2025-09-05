import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt
from PIL import Image
import google.generativai as genai

# ===================================================================
# Налаштування AI ключа та моделі
# ===================================================================
AI_MODEL_NAME = 'gemini-1.5-flash-latest'

try:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
except Exception as e:
    st.error("Помилка конфігурації Google AI. Переконайтеся, що ви створили файл .streamlit/secrets.toml і додали туди ваш GOOGLE_API_KEY.")
    st.stop()

# ===================================================================
# КЕШОВАНІ ФУНКЦІЇ ДЛЯ РОБОТИ З AI (ЕКОНОМІЯ ЧАСУ ТА ГРОШЕЙ)
# ===================================================================

@st.cache_data
def ocr_with_gemini_vision(img_bytes: bytes, filename: str):
    """Розпізнає текст на зображенні за допомогою Gemini Vision."""
    img = Image.open(BytesIO(img_bytes))
    try:
        model = genai.GenerativeModel(AI_MODEL_NAME)
        prompt = "Ти — експертна система OCR. Максимально точно розпізнай та транскрибуй весь український текст на цьому зображенні. Поверни лише чистий текст."
        response = model.generate_content([prompt, img])
        return response.text
    except Exception as e:
        st.error(f"Помилка розпізнавання ({filename}): {e}")
        return None

@st.cache_data
def process_text_with_gemini(raw_text: str, style: str, filename: str):
    """Обробляє текст згідно з обраним стилем."""
    if not raw_text:
        return ""
    
    prompts = {
        "Літературне редагування": f"""
            Ти — досвідчений літературний редактор. Твоє завдання — взяти сирий текст і перетворити його на повноцінну, чисту та читабельну статтю.
            Інструкції:
            1. Придумай влучний заголовок.
            2. Ретельно виправ всі помилки (орфографічні, граматичні).
            3. Розбий текст на логічні абзаци.
            4. Видали будь-які артефакти OCR.
            5. Поверни ТІЛЬКИ відформатовану статтю із заголовком. Не пиши жодних коментарів.
            Ось сирий текст: --- {raw_text} ---
        """,
        "Тільки корекція помилок": f"""
            Ти — уважний коректор. Твоє завдання — виправити помилки в тексті, не змінюючи його структуру.
            Інструкції:
            1. Виправ орфографічні, граматичні та пунктуаційні помилки.
            2. НЕ змінюй розбиття на абзаци і НЕ додавай заголовок.
            3. Поверни тільки виправлений текст.
            Ось сирий текст: --- {raw_text} ---
        """,
        "Стислий переказ (тези)": f"""
            Ти — аналітик. Прочитай цей текст і напиши його стислий переказ у вигляді тез (ключові думки).
            Поверни тільки тези.
            Ось сирий текст: --- {raw_text} ---
        """
    }
    
    try:
        model = genai.GenerativeModel(AI_MODEL_NAME)
        response = model.generate_content(prompts[style])
        return response.text
    except Exception as e:
        st.warning(f"Помилка обробки ({filename}): {e}")
        return raw_text

@st.cache_data
def generate_summary_page(_processed_articles: tuple):
    """Створює титульну сторінку-зміст для всього документа."""
    full_text = "\n\n--- НОВА СТАТТЯ ---\n\n".join(_processed_articles)
    try:
        model = genai.GenerativeModel(AI_MODEL_NAME)
        prompt = f"""
        Ти — науковий асистент. Проаналізуй набір статей, наведений нижче.
        Твоє завдання — створити титульну сторінку та зміст для наукової роботи.
        
        Структура відповіді:
        1.  **Загальний заголовок**: Придумай загальний заголовок для всього збірника статей.
        2.  **Зміст**: Для КОЖНОЇ статті напиши:
            - Її порядковий номер та назву.
            - Короткий переказ (1-2 речення).
            - Ключові особи, організації та локації, що в ній згадуються.
        3.  **Загальні ключові слова**: В кінці напиши 5-7 ключових слів для всього документа.
        
        Поверни тільки цю сторінку, без зайвих коментарів.
        Ось статті:
        ---
        {full_text}
        ---
        """
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Помилка створення змісту: {e}")
        return "Не вдалося створити автоматичний зміст."

# =======================
# Функції для роботи з DOCX
# =======================
def save_full_document_to_docx(summary_page: str, articles: list):
    """Збирає титульну сторінку та обрані статті в один DOCX файл."""
    doc = Document()
    doc.add_heading('Автоматичний зміст та аналіз документа', level=1)
    doc.add_paragraph(summary_page)
    doc.add_page_break()
    
    for article_text in articles:
        doc.add_paragraph(article_text)
        doc.add_page_break()
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===================================================================
# ІНТЕРФЕЙС ПРОГРАМИ (Streamlit UI)
# ===================================================================
st.set_page_config(page_title="Архівний Асистент PRO", layout="wide")
st.title("🗂️ Архівний Асистент PRO")

if "files_data" not in st.session_state:
    st.session_state.files_data = {}

st.header("Крок 1: Налаштуйте обробку та завантажте файли")

col1, col2 = st.columns(2)
with col1:
    processing_style = st.selectbox(
        "Оберіть стиль редагування AI:",
        ("Літературне редагування", "Тільки корекція помилок", "Стислий переказ (тези)")
    )
with col2:
    uploaded_files = st.file_uploader(
        "Завантажте одразу декілька зображень:",
        type=["jpg", "png"],
        accept_multiple_files=True
    )

if uploaded_files:
    if st.button("🚀 Розпочати обробку всіх файлів", type="primary"):
        st.session_state.files_data = {}
        with st.spinner("AI працює... Це може зайняти деякий час."):
            for i, uploaded_file in enumerate(uploaded_files):
                # =================================================
                # ОСЬ ТУТ БУЛО ВИПРАВЛЕННЯ: .id -> .file_id
                # =================================================
                file_id = uploaded_file.file_id
                
                img_bytes = uploaded_file.getvalue()
                
                raw_text = ocr_with_gemini_vision(img_bytes, uploaded_file.name)
                processed_text = process_text_with_gemini(raw_text, processing_style, uploaded_file.name)
                
                st.session_state.files_data[file_id] = {
                    "name": uploaded_file.name,
                    "raw": raw_text,
                    "processed": processed_text,
                    "selected": True
                }
        st.success("Обробку всіх файлів завершено!")

if st.session_state.files_data:
    st.header("Крок 2: Перегляньте результати та оберіть статті для документа")
    
    for file_id, data in st.session_state.files_data.items():
        with st.expander(f"{data['name']}", expanded=False):
            data['selected'] = st.checkbox("Додати цю статтю до фінального документа", value=data['selected'], key=f"cb_{file_id}")
            st.text_area("Результат обробки", data['processed'], height=200, key=f"txt_{file_id}")

    st.header("Крок 3: Створіть та завантажте фінальний документ")
    
    selected_articles = [data['processed'] for data in st.session_state.files_data.values() if data['selected']]
    
    if not selected_articles:
        st.warning("Ви не обрали жодної статті для додавання в документ.")
    else:
        st.info(f"Буде створено документ з {len(selected_articles)} обраних статей.")
        if st.button("📝 Створити фінальний документ з титульною сторінкою"):
            with st.spinner("AI генерує зміст..."):
                summary_page_text = generate_summary_page(tuple(selected_articles))
            
            st.subheader("Автоматично створений зміст:")
            st.markdown(summary_page_text)

            docx_buffer = save_full_document_to_docx(summary_page_text, selected_articles)
            
            st.download_button(
                label="📥 Завантажити готовий документ (.docx)",
                data=docx_buffer,
                file_name="PRO_Archive_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

